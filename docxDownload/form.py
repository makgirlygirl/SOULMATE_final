#%%

## import
from docx import Document
from docx.shared import Pt
import os
#%%
# data_dir = '/home/a1930008/k_sat_pre_hong/test_all_Q_category'
# data_dir='/home/a1930008/k_sat_pre_hong/newk_sat/test_all_Q_category/result'
# data_dir='/home/a1930008/k_sat_pre_hong/newk_sat/old_test_all_Q_category/old_result'
data_dir='/home/a1930008/k_sat_pre_hong/newk_sat/test_all_Q_category/realfinal'

file_list =  os.listdir(data_dir)
file_path=[]
for i in file_list:
    if i.split('.')[-1]=='txt':
        file_path.append(os.path.join(data_dir, i))
del file_list

#%%

input=[]
for i in file_path:
    f = open(i)
    ff=f.readlines()
    for j in ff:
        if j.startswith('{'):
            input.append(eval(j))
#%%
class makeDocx:
    def __init__(self):
        self.doc = Document('/home/a1930008/docxDownload/form_a4.docx')
        style = self.doc.styles['Normal']
        font=style.font
        font.name='Times New Roman'
        font.size=Pt(9)
        self.Qnumber=1
        self.ansAll=[]

    def addQ(self,dict):
        doc=self.doc
        doc.add_paragraph(str(self.Qnumber)+". "+dict["question"]+'\n') ##제목
        doc.add_paragraph(" "+dict["new_passage"]+'\n') ##본문

        ans=dict["answer"]  ## 정답 번호 
        self.ansAll.append(ans)

        doc.add_paragraph('① '+str(dict["e1"]))
        doc.add_paragraph('② '+str(dict["e2"]))
        doc.add_paragraph('③ '+str(dict["e3"]))
        doc.add_paragraph('④ '+str(dict["e4"]))
        doc.add_paragraph('⑤ '+str(dict["e5"]))      
        doc.add_paragraph("\n")
        self.Qnumber+=1

    def recur(self,dict_list):
        doc=self.doc
        for i in range(len(dict_list)):
            self.addQ(dict_list[i])

    def addA(self):
        doc=self.doc
        doc.add_page_break()
        doc.add_paragraph("<Answer>")
        for i in range(len(self.ansAll)):
            doc.add_paragraph(str(i+1)+') '+str(self.ansAll[i]))

    def saveDocx(self,dict_list,docName="/home/a1930008/docxDownload/download.docx"):
        doc=self.doc
        self.recur(dict_list)
        self.addA()
        doc.save(docName)
# %%
# d=makeDocx()
# d.__init__()
# d.saveDocx(input)


# %%
